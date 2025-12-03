import streamlit as st
import os
import requests
import docx
from datetime import datetime
import io

st.session_state.ended = False

# Custom CSS for styling
st.markdown("""
    <style>
    .title {
        font-size:30px !important;
        #color:#FF5733;
        color:#2E86C1;     
        text-align:center;
    }
    .subtitle {
        font-size:20px !important;
        color:#FF5733;
        text-align:left;
    }
    .highlight {
        color:#28B463;  
        font-weight:bold;
    }
    </style>
""", unsafe_allow_html=True)


st.markdown("<p class='title'>Risk Analysis Report Generator</p>", unsafe_allow_html=True)

sentence = """
This module generates a comprehensive Cyber Risk and Resiliency Report for industrial systems, using the results of a prior sensitivity analysis. 
 It prompts the user for an API key and business sector. The tool constructs a detailed prompt and sends it to an AI model to produce a report tailored for senior leadership. 
 The generated report is displayed in the tool and can be downloaded as a formatted Word document. 
 The report covers executive summary, risk analysis, sensitivity outcomes, and actionable recommendations, all based on the system’s dependency structure and risk metrics. 
"""

st.markdown(f"<p class='highlight'>{sentence}</p>", unsafe_allow_html=True)


if 'ended' not in st.session_state:
      st.session_state.ended = False

if st.session_state.ended:
      st.info("The Module produced the Reports. It is the last action to be performed. Please close this tab and return to the RISKED Menu tab.")
      st.stop()


# --- CONFIGURATION ---
DEFAULT_MODEL = "llama-3.1-8b-instant"
MODEL = DEFAULT_MODEL
required_files = [
    "one_point_sensitivity.csv",
    "data_model_with_posterior.csv"
]
# Hardcoded values
sector = "Manufacturing"


# --- Ask user for API key in sidebar ---
st.sidebar.header("Groq/Llama API Key")
api_key = st.sidebar.text_input(
    "Enter your Groq/Llama API Key:",
    type="password",
    help="Paste your Groq/Llama API key here. It will be used for all API requests."
)

if not api_key:
    st.warning("Please enter your Groq/Llama API key in the sidebar to continue.")
    st.stop()


# In the sidebar, ask for company name
sector = st.sidebar.text_input(
    "Enter your organisation's Sector:",
    help="Type your organisation's Sector so as to focus the analysis."
)

if not sector:
    st.warning("Please enter your organisation's Sector in the sidebar to continue.")
    st.stop()


PROMPT_TEMPLATE = """

Business Sector: {sector}

You are an AI expert in cybersecurity risk analysis for industrial systems (e.g., SCADA, ICS). Your expertise includes probabilistic modelling, sensitivity analysis, and interpreting how failures or attacks propagate through complex dependencies. You are familiar with standards such as MITRE ATT&CK for ICS, NIST SP 800-82, and ISA/IEC 62443.

You are provided with two CSV files:
- one_point_sensitivity.csv
- data_model_with_posterior.csv

**File Descriptions:**
- *one_point_sensitivity.csv*: Results of single-node sensitivity analysis. The key column, `prob_given_node0`, shows the marginal probability of the root node if a given leaf node fails completely (0% reliability).
- *data_model_with_posterior.csv*: Contains the Bayesian Network structure, node dependencies, and baseline probabilities, including the root node.

**Task:**  
Generate a comprehensive Cyber Risk and Resiliency Report for senior leadership (CISO, CRO, Asset Owners). The report should be accessible to non-technical readers but analytically rigorous. Use your own words and avoid copying phrases from this prompt.
In producing your report, please mutiply all probabilities by 100 to convert them into percentages with two decimal places.
**Report Structure:**

1. **Executive Summary**
   - Briefly describe the system and its dependency model.
   - Explain the meaning of the "root node" (e.g., overall system availability).
   - State the current marginal probability of the root node and interpret its business impact.
   - Summarise key system characteristics (dependency depth, complexity, coupling).
   - Highlight the most critical leaf nodes and their impact.

2. **Probabilistic Risk Analysis**
   - Identify 2 to 3 plausible failure scenarios based on sensitivity scores and dependencies.
   - For each scenario, explain:
     - What can go wrong (business operations impact).
     - How likely it is (qualitative and quantitative assessment).
     - Potential consequences (business, safety, regulatory, financial), referencing relevant frameworks.

3. **Sensitivity Analysis Outcomes**
   - Use one_point_sensitivity.csv to identify the top 5 most critical leaf nodes (lowest Probability Sensitivity Scores).
   - For each node, explain:
     - How its failure affects the root node’s risk.
     - How root node reliability improves if the node is fully reliable.
     - Provide formatted Probability Sensitivity Scores (e.g., “Root node reliability improves to 12.34% if Node X is fully reliable”).
   - Use data_model_with_posterior.csv to discuss interdependencies and cascading effects.

4. **Recommendations & Strategic Reminders**
   - Suggest targeted mitigation strategies for the top 5 critical nodes, referencing standards where appropriate.
   - Conclude with two paragraphs on essential security hygiene and foundational practices, referencing relevant frameworks.

**Formatting Instructions:**
- Format all probabilities and sensitivity scores as percentages with two decimal places (e.g., 21.34%).
- Structure the report with clear headings and concise paragraphs.

**Data Provided:**
--- one_point_sensitivity.csv ---
{one_point_sensitivity}
--- data_model_with_posterior.csv ---
{data_model_with_posterior}

"""

def read_file_as_text(filename):
    try:
        with open(filename, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return None

def call_groq(prompt, api_key, model):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI expert in cybersecurity risk analysis for complex industrial systems like SCADA and ICS. "
                    "Your specialty is using Inference (Variable Elimination), Sensitivity Analysis & Parallelisation, probabilistic models to understand how failures and attacks can propagate through interdependent components. "
                    "Provide detailed, actionable analysis and recommendations. Be comprehensive and professional. Let your ourput be professionally formatted for the audience of senior leadership (CISO, CRO, Asset Owners)."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.6,
        "max_tokens": 4000
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Unexpected Error: {str(e)}"

def generate_word_report(report, sector):
    doc = docx.Document()
    doc.add_heading('Executive Risk & Resiliency Report', 0)
    doc.add_paragraph(f"Business Sector: {sector}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    lines = report.split('\n')
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == "---":
            doc.add_paragraph("")
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip()[2:-2]).bold = True
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("Generate and Download Report"):
    # Step 1: Check files
    missing_files = [f for f in required_files if not os.path.isfile(f)]
    
    if missing_files:
        st.error(f"The required file(s) to perform this process are missing."
        "You may have missed a step in the analysis sequence. "
        "Please complete the previous analysis steps to generate these files before proceeding.")
        st.stop()
    
    else:
        file_texts = {filename: read_file_as_text(filename) or "" for filename in required_files}
        prompt = PROMPT_TEMPLATE.format(
            sector=sector,
            one_point_sensitivity=file_texts["one_point_sensitivity.csv"],
            data_model_with_posterior=file_texts["data_model_with_posterior.csv"]
        )
        report = call_groq(prompt, api_key, MODEL)
        if report.startswith("API Error") or report.startswith("Unexpected Error"):
            st.error(f"Please provide a valid API key to proceed. You may go to Groq/Llama to generate a free API Key, if you do not have one.")
            st.stop()
            #st.error(report)
        else:
            # Display the report on the screen
            st.subheader("Generated Risk Analysis Report")
            st.markdown(report)
            # Provide download button for Word doc
            word_buffer = generate_word_report(report, sector)
            st.download_button(
                label="Download Risk Analysis Report (Word)",
                data=word_buffer,
                file_name=f"Risk_Resiliency_Report_{sector.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
st.caption("This is a PoC and it is powered by Groq and Llama 3. In the full version, this is powered by our purpose-built genAI for Cyber risk identification and assessement.")
st.markdown("---")    
if st.button("End Module"):
      st.session_state.ended = True
      st.rerun()
      