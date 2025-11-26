import streamlit as st
import pandas as pd
import requests
from datetime import datetime
import docx
import io
import os

st.session_state.ended = False
# if 'ended' not in st.session_state:
#       st.session_state.ended = False

# if st.session_state.ended:
#       st.info("The Cyber Risk Security Profile Report Download is completed.")
#       st.stop()


# Custom CSS for styling
st.markdown("""
    <style>
    .title {
        font-size:35px !important;
        #color:#FF5733;
        color:#2E86C1;     
        text-align:left;
    }
    .subtitle {
        font-size:30px !important;
        color:#FF5733;
        text-align:centre;
    }
    .highlight {
        font-size:20px !important;
        color:#28B463;  
        font-weight:bold;
    }
    </style>
""", unsafe_allow_html=True)
st.session_state.ended = False

st.markdown("<p class='title'>Cyber Risk Profile Report Generator</p>", unsafe_allow_html=True)


# ========================
# CONFIGURATION
# ========================

#COMPANY_NAME = "Demo - CyberRisk Resilience & Security Methods INC"
AI_PROVIDER = "groq"
MODEL = "llama-3.1-8b-instant"
CSV_FILE = "Security_Profile_Responses.csv"


# --- NEW: Ask user for API key in sidebar ---

  
st.sidebar.header("Groq/Llama API Key")
API_KEY = st.sidebar.text_input(
    "Enter your Groq/Llama API Key:",
    type="password",
    help="Paste your Groq/Llama API key here. It will be used for all API requests."
)

if not API_KEY:
    st.warning("Please enter your Groq/Llama API key in the sidebar to continue.")
    st.stop()


# In the sidebar, ask for company name
COMPANY_NAME = st.sidebar.text_input(
    "Enter your Company Name:",
    help="Type your organisation's name as you want it to appear in the report."
)

if not COMPANY_NAME:
    st.warning("Please enter your company name in the sidebar to continue.")
    st.stop()



# ========================
# API CALL
# ========================
def call_groq(prompt, api_key, model):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a cybersecurity expert specializing in ICS/OT security frameworks and risk assessment. Provide detailed, actionable analysis and recommendations. Be comprehensive and professional in your analysis."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.5,
        "max_tokens": 4000
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            return f"❌ API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"❌ Unexpected Error: {str(e)}"

def generate_cybersecurity_analysis(csv_content, api_key, model):
    prompt = f"""
ANALYZE THIS CYBERSECURITY PROFILE DATA AND GENERATE A COMPREHENSIVE REPORT:
[CSV DATA START]
{csv_content}
[CSV DATA END]
I have used 12 cybersecurity frameworks and standards to curate the statements in this csv file. A user has provided responses to each of the statements. The 12 standards and framework are: ISA/IEC 62443 Series, NIST SP 800-82, ISO/IEC 27001, NIST Cybersecurity Framework (CSF), Cyber Essentials Plus, CIS Critical Security Controls (CIS Controls), ENISA ICS Security Guidelines, CPNI Guidance, NIS2 Directive & EU Cyber Resilience Act, ISO/IEC 27019, NIST SP 800-53, ISA/IEC 61511. Now, could you use the responses from the user, as shown in "Response" and "Evidence" columns in the CSV to determine and generate a holistic Cybersecurty profile for the user. Please note, "Strongly agreed" means they are fully compliant and "Strongly Disagree" means they are not compliant at all. The "Evidence" value will only be relevant if the user picked "Strongly Agree" or "Agreed". A "Not Applicable" response means the statement is not applicable to the user. The value in the "Weight" column, is used to determine how important the statement requirement is and how high the risk is if not compliant with high the risk is, where 4 is the highest. Please also include gap analysis and what the user could do to improve the cybersecurity posture/profile.
Please structure your response as follows:
# Profile Summary
[Overall assessment of cybersecurity posture]
**Overall Cybersecurity Maturity Score:** [Percentage score with explanation]
---
## Detailed Gap Analysis & Improvement Plan
### 1. System & Organizational Resilience
* **Critical Gaps:**
[List critical gaps]
* **Recommendations:**
[List specific recommendations]
### 2. Technology & Technical Controls
* **Critical Gaps:**
[List critical gaps]
* **Recommendations:**
[List specific recommendations]
### 3. Process & Governance
* **Critical Gaps:**
[List critical gaps]
* **Recommendations:**
[List specific recommendations]
### 4. People & Awareness
* **Critical Gaps:**
[List critical gaps]
* **Recommendations:**
[List specific recommendations]
## Conclusion and Strategic Next Steps
**Immediate Priorities (Next 3-6 Months):**
1. [Priority 1]
2. [Priority 2]
3. [Priority 3]
**Long-term Strategic Goals (6-18 Months):**
1. [Goal 1]
2. [Goal 2]
"""
    return call_groq(prompt, api_key, model)

def create_word_document(analysis_text, response_df):
    doc = docx.Document()
    doc.add_heading('DEMO - Cyber Risk Profile Analysis', 0)
    doc.add_paragraph(f"Organization: {COMPANY_NAME}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    lines = analysis_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == '---':
            doc.add_paragraph()
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.strip()[2:-2]).bold = True
        elif line.strip():
            doc.add_paragraph(line)
        i += 1
    doc.add_paragraph()
    doc.add_heading('Assessment Methodology', level=2)
    doc.add_paragraph("This analysis is based on 12 industry-standard cybersecurity frameworks and standards:")
    frameworks = [
        "ISA/IEC 62443 Series", "NIST SP 800-82", "ISO/IEC 27001", "NIST Cybersecurity Framework (CSF)",
        "Cyber Essentials Plus", "CIS Critical Security Controls (CIS Controls)", "ENISA ICS Security Guidelines",
        "CPNI Guidance", "NIS2 Directive & EU Cyber Resilience Act", "ISO/IEC 27019", "NIST SP 800-53", "ISA/IEC 61511"
    ]
    for framework in frameworks:
        doc.add_paragraph(framework, style='List Bullet')
    doc.add_paragraph()
    doc.add_heading('Raw Response Data', level=2)
    table = doc.add_table(rows=1, cols=len(response_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(response_df.columns):
        hdr_cells[i].text = str(column)
    for index, row in response_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    return doc

# ========================
# STREAMLIT APP
# ========================

#st.markdown("<p class='title'>Cybersecurity Risk Profile Report Generator</p>", unsafe_allow_html=True)
st.markdown("<p class='subtitle'>Cyber Risk Profile Report</p>", unsafe_allow_html=True)

# st.set_page_config(page_title="Cyber-risk Security Profile Report", layout="wide")
# st.title("Cybersecurity Profile Report")
coyx = (f"Organization:  {COMPANY_NAME}")
st.markdown(f"<p class='highlight'>{coyx}</p>", unsafe_allow_html=True)

# Automatically process the report
if os.path.exists(CSV_FILE):
    response_df = pd.read_csv(CSV_FILE)
    required_columns = ['Category', 'Security Profile Statement', 'Weight', 'Response', 'Evidence']
    missing_columns = [col for col in required_columns if col not in response_df.columns]
    if missing_columns:
        st.error(f"❌ Error: Missing required columns in CSV: {missing_columns}")
    else:
        csv_content = response_df.to_csv(index=False)
        analysis_result = generate_cybersecurity_analysis(csv_content, API_KEY, MODEL)
        if analysis_result.startswith("❌"):
            st.error(analysis_result)
        else:
            #st.header("Cybersecurity Analysis Report")
            st.markdown(analysis_result)
            # Generate Word document
            doc = create_word_document(analysis_result, response_df)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            safe_company_name = COMPANY_NAME.replace(' ', '_').replace('/', '_')
            output_filename = f"cybersecurity_analysis_report_{safe_company_name}_{timestamp}.docx"
            st.download_button(
                label="Download Report (Word doc)",
                data=doc_io,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.error(f"❌ Error: CSV file '{CSV_FILE}' not found in the app directory.")

st.caption("This is a PoC and it is powered by Groq and Llama 3. In the full version, this is powered by our purpose-built genAI for Cyber risk identification and assessement.")

st.markdown("---")    
if st.button("End Module"):
      st.session_state.ended = True
      st.info("The Cyber Risk Security Profile Report Download is completed.")
      #st.rerun()